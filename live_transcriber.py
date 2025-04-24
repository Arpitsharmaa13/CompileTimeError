import os
import nltk
import re
import asyncio
import json
import google.generativeai as genai
import sounddevice as sd
import numpy as np
from scipy.io.wavfile import write
from datetime import datetime, timedelta
from deepgram import Deepgram
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from nltk.tokenize import sent_tokenize
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import shutil
import time

# NLTK Setup
nltk.download('punkt', quiet=True)

# API Configurations
DEEPGRAM_API_KEY = os.getenv("DEEPGRAM_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

# Google Calendar Scope
SCOPES = ['https://www.googleapis.com/auth/calendar']

# Function to authenticate Google Calendar API
def authenticate_google_account():
    creds = None
    if os.path.exists('token.json'):
        try:
            creds = Credentials.from_authorized_user_file('token.json', SCOPES)
        except Exception as e:
            print(f"Error loading token.json: {e}")
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                print("Google Calendar token refreshed")
            except Exception as e:
                print(f"Error refreshing token: {e}")
                creds = None
        if not creds:
            credentials_path = 'post_meeting/config/credentials.json'
            if not os.path.exists(credentials_path):
                raise FileNotFoundError(f"Google Calendar credentials file not found at {credentials_path}")
            try:
                flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
                creds = flow.run_local_server(port=0)
                print("Google Calendar authenticated")
            except Exception as e:
                print(f"Error authenticating Google Calendar: {e}")
                return None
        
        try:
            with open('token.json', 'w') as token:
                token.write(creds.to_json())
        except Exception as e:
            print(f"Error saving token.json: {e}")
    
    try:
        service = build('calendar', 'v3', credentials=creds)
        return service
    except Exception as e:
        print(f"Error building Google Calendar service: {e}")
        return None

# Function to extract key points from text
def extract_key_points(text):
    sentences = sent_tokenize(text)
    key_points = []
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in ['by', 'next', 'decision', 'proposal']):
            key_points.append(sentence.strip())
    if key_points:
        print("\nKey Points:")
        for point in key_points:
            print(f"- {point}")
    else:
        print("\nKey Points: None found.")
    return key_points

# Function to remove suffixes (st, nd, rd, th) from dates
def clean_date_suffix(date_str):
    return re.sub(r'(st|nd|rd|th)', '', date_str)

# Function to extract dates from text
def extract_dates(text):
    current_year = datetime.now().year
    current_date = datetime.now()
    valid_months = [
        'January', 'February', 'March', 'April', 'May', 'June',
        'July', 'August', 'September', 'October', 'November', 'December'
    ]
    
    date_patterns = [
        r'\b(\w+ \d{1,2},? \d{4})\b',
        r'\b(\d{1,2} \w+ \d{4})\b',
        r'\b(\w+ \d{1,2}(?:st|nd|rd|th)?)\b',
        r'\b(\d{1,2}(?:st|nd|rd|th)? \w+)\b',
        r'\b(next (?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday))\b',
        r'\b(tomorrow)\b',
        r'\b(next week)\b'
    ]
    
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        dates.extend(matches)
    
    cleaned_dates = []
    for date in dates:
        date = clean_date_suffix(date).strip()
        print(f"Detected date: {date}")
        
        try:
            if date.lower() == 'tomorrow':
                parsed_date = current_date + timedelta(days=1)
                cleaned_dates.append(parsed_date.strftime("%B %d %Y"))
                continue
            elif date.lower() == 'next week':
                parsed_date = current_date + timedelta(days=7)
                cleaned_dates.append(parsed_date.strftime("%B %d %Y"))
                continue
            elif date.lower().startswith('next '):
                day = date[5:].capitalize()
                days_until = (datetime.strptime(day, "%A").weekday() - current_date.weekday() + 7) % 7
                if days_until == 0:
                    days_until = 7
                parsed_date = current_date + timedelta(days=days_until)
                cleaned_dates.append(parsed_date.strftime("%B %d %Y"))
                continue
            
            parts = date.split()
            if len(parts) >= 2:
                month, day = None, None
                if parts[0] in valid_months:
                    month = parts[0].capitalize()
                    day = parts[1]
                elif parts[1] in valid_months:
                    month = parts[1].capitalize()
                    day = parts[0]
                
                if month and day:
                    year = parts[2] if len(parts) > 2 and parts[2].isdigit() else current_year
                    date_str = f"{month} {day} {year}"
                    parsed_date = datetime.strptime(date_str, "%B %d %Y")
                    cleaned_dates.append(parsed_date.strftime("%B %d %Y"))
        except Exception as e:
            print(f"Error parsing date '{date}': {e}")
    
    if cleaned_dates:
        print(f"Parsed dates: {cleaned_dates}")
    else:
        print("No valid dates parsed")
    return cleaned_dates

# Function to infer event title from transcript
def infer_event_title(transcript):
    keywords = ['meeting', 'follow-up', 'proposal', 'review', 'discussion']
    for keyword in keywords:
        if keyword in transcript.lower():
            return f"{keyword.capitalize()} on {{date}}"
    return "Meeting on {date}"

# Function to sync event with Google Calendar
def sync_event(start_time, end_time, event_title):
    service = authenticate_google_account()
    if not service:
        print("Cannot sync event: Google Calendar service unavailable")
        return
    
    event = {
        'summary': event_title,
        'start': {
            'dateTime': start_time,
            'timeZone': 'America/Los_Angeles',
        },
        'end': {
            'dateTime': end_time,
            'timeZone': 'America/Los_Angeles',
        },
    }
    try:
        print(f"Sending event to Google Calendar: {event_title}")
        event_result = service.events().insert(calendarId='primary', body=event).execute()
        print(f"Event synced: {event_result['summary']} at {event_result['start']['dateTime']}")
    except Exception as e:
        print(f"Failed to sync event '{event_title}': {e}")

# Function to query Gemini API with retry
def query_gemini(prompt, max_tokens=150, retries=3):
    for attempt in range(retries):
        try:
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(
                f"You are an expert in marketing and sales, specializing in persuasive communication, customer engagement, and strategic business solutions. Provide concise, professional, and customer-focused answers that align with marketing and sales goals. {prompt}",
                generation_config={"max_output_tokens": max_tokens}
            )
            answer = response.text.strip()
            with open("gemini_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(f"Prompt: {prompt}\nAnswer: {answer}\n\n")
            return answer
        except Exception as e:
            print(f"Gemini API attempt {attempt + 1}/{retries} failed: {e}")
            if attempt < retries - 1:
                asyncio.sleep(2)
            else:
                with open("gemini_log.txt", "a", encoding="utf-8") as log_file:
                    log_file.write(f"Prompt: {prompt}\nError: {e}\n\n")
                return f"Error querying Gemini API: {e}"

# Function to identify and answer queries
def answer_queries(transcript, query_answers):
    queries = re.findall(r'(?:\s|^)([A-Z][^\.]*\?)(?:\s|$)', transcript)
    if not queries:
        print("\nNo queries detected")
        return

    print("\nQuery Answers (Marketing Specialist):")
    for query in queries:
        query = query.strip()
        print(f"Detected query: {query}")
        try:
            prompt = f"Answer this query as a marketing specialist: {query}"
            answer = query_gemini(prompt, max_tokens=150)
            query_answers.append((query, answer))
            print(f"\nQuery: {query}\nAnswer: {answer}")
        except Exception as e:
            print(f"Error for '{query}': {e}")

# Function to create Word document
def create_word_document(transcript, key_points, query_answers, dates):
    doc = Document()
    
    # Title
    title = doc.add_heading('Meeting Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Full Transcript
    doc.add_heading('Full Transcript', level=1)
    doc.add_paragraph(transcript if transcript else "No transcript generated.", style='Normal').runs[0].font.size = Pt(12)
    
    # Key Points
    doc.add_heading('Key Points', level=1)
    if key_points:
        for point in key_points:
            doc.add_paragraph(point, style='ListBullet').runs[0].font.size = Pt(12)
    else:
        doc.add_paragraph('No key points identified.', style='ListBullet')
    
    # Important Dates and Deadlines
    doc.add_heading('Important Dates and Deadlines', level=1)
    if dates:
        for date in dates:
            doc.add_paragraph(f"Event scheduled for {date}", style='ListBullet').runs[0].font.size = Pt(12)
    else:
        doc.add_paragraph('No dates or deadlines identified.', style='ListBullet')
    
    # Query Answers
    doc.add_heading('Query Answers (Marketing Specialist)', level=1)
    if query_answers:
        for query, answer in query_answers:
            p = doc.add_paragraph()
            p.add_run(f"Query: {query}").bold = True
            p.add_run(f"\nAnswer: {answer}").font.size = Pt(12)
    else:
        doc.add_paragraph('No queries identified.', style='ListBullet')
    
    # Save document with retry logic
    output_file = 'D:\\meeting_summary.docx'
    output_dir = os.path.dirname(output_file) or '.'
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            # Check if directory is writable
            if not os.access(output_dir, os.W_OK):
                raise PermissionError(f"No write permission in directory: {output_dir}")
            
            # Create temp file in same directory
            temp_dir = output_dir
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=temp_dir)
            temp_path = temp_file.name
            temp_file.close()
            
            # Save document to temp file
            doc.save(temp_path)
            
            # Delete existing output file if it exists
            if os.path.exists(output_file):
                try:
                    os.remove(output_file)
                except Exception as e:
                    print(f"Warning: Could not delete existing {output_file}: {e}")
            
            # Move temp file to final destination
            shutil.move(temp_path, output_file)
            print(f"\nWord document saved to {output_file}")
            return
        
        except Exception as e:
            error_msg = f"Error saving Word document (attempt {attempt + 1}/{max_retries}): {str(e)}"
            print(error_msg)
            try:
                with open("word_doc_log.txt", "a", encoding="utf-8") as log_file:
                    log_file.write(f"{error_msg}\n")
            except Exception:
                print("Failed to write to word_doc_log.txt due to encoding issues")
            if attempt < max_retries - 1:
                time.sleep(1)  # Wait before retrying
            else:
                # Fallback to user's home directory
                try:
                    fallback_path = os.path.join(os.path.expanduser("~"), "meeting_summary_fallback.docx")
                    doc.save(fallback_path)
                    print(f"Fallback: Word document saved to {fallback_path}")
                except Exception as fallback_e:
                    error_msg = f"Fallback save failed: {str(fallback_e)}"
                    print(error_msg)
                    try:
                        with open("word_doc_log.txt", "a", encoding="utf-8") as log_file:
                            log_file.write(f"{error_msg}\n")
                    except Exception:
                        print("Failed to write to word_doc_log.txt due to encoding issues")

# Async function to listen for Enter key to stop recording
async def input_listener(stop_event):
    loop = asyncio.get_event_loop()
    def check_input():
        input()  # Blocks until Enter is pressed
        stop_event.set()
    await loop.run_in_executor(None, check_input)

# Async function to record and transcribe audio
async def record_and_transcribe(audio_file="meeting_recording.wav", sample_rate=16000):
    if not DEEPGRAM_API_KEY:
        print("Error: DEEPGRAM_API_KEY environment variable not set.")
        return "", [], [], []

    print("Press Enter to start recording audio...")
    input()
    print("Recording audio... (Press Enter to stop)")

    recording = []
    transcript = ""
    key_points = []
    query_answers = []
    dates = []
    stop_event = asyncio.Event()

    def audio_callback(indata, frames, time, status):
        if status:
            print(f"Audio callback status: {status}")
        recording.append(indata.copy())

    try:
        # Start input listener task
        input_task = asyncio.create_task(input_listener(stop_event))
        
        # Record audio
        with sd.InputStream(samplerate=sample_rate, channels=1, callback=audio_callback):
            await stop_event.wait()
        
        # Cancel input task
        input_task.cancel()
        try:
            await input_task
        except asyncio.CancelledError:
            pass

        print("\nRecording stopped")
        if recording:
            try:
                recording = np.concatenate(recording, axis=0)
                write(audio_file, sample_rate, recording)
                print(f"\nAudio saved to {audio_file}")
            except Exception as e:
                print(f"Error saving audio: {e}")
                return "", [], [], []
        else:
            print("No audio recorded.")
            return "", [], [], []

        # Transcribe WAV file
        if os.path.exists(audio_file) and os.path.getsize(audio_file) > 0:
            print(f"\nTranscribing audio file: {audio_file}")
            dg_client = Deepgram(DEEPGRAM_API_KEY)
            
            for attempt in range(3):
                try:
                    print(f"Sending audio to Deepgram (attempt {attempt + 1}/3)...")
                    with open(audio_file, 'rb') as audio:
                        response = dg_client.transcription.sync_prerecorded(
                            {'buffer': audio, 'mimetype': 'audio/wav'},
                            {'model': 'nova-2', 'language': 'en-US', 'smart_format': True}
                        )
                    transcript = response['results']['channels'][0]['alternatives'][0]['transcript']
                    if transcript:
                        print(f"\nTranscript: {transcript}")
                        with open("deepgram_log.txt", "a", encoding="utf-8") as log_file:
                            log_file.write(f"{json.dumps(response)}\n")
                        break
                    else:
                        print(f"Empty transcript received on attempt {attempt + 1}")
                except Exception as e:
                    print(f"Transcription attempt {attempt + 1}/3 failed: {e}")
                    with open("deepgram_log.txt", "a", encoding="utf-8") as log_file:
                        log_file.write(f"Transcription error (attempt {attempt + 1}): {e}\n")
                    if attempt < 2:
                        await asyncio.sleep(2)
                    else:
                        print("All transcription attempts failed.")
                        return "", [], [], []

            if transcript:
                # Save transcript
                try:
                    with open("transcript.txt", "w", encoding="utf-8") as file:
                        file.write(transcript)
                    print("Transcript saved to transcript.txt")
                except Exception as e:
                    print(f"Error saving transcript: {e}")

                # Process transcript
                key_points = extract_key_points(transcript)
                answer_queries(transcript, query_answers)
                
                # Extract dates and sync to Google Calendar
                dates = extract_dates(transcript)
                if dates:
                    print("\nSyncing Dates to Google Calendar...")
                    event_title_template = infer_event_title(transcript)
                    for date in dates:
                        try:
                            parsed_date = datetime.strptime(date, "%B %d %Y")
                            event_start_time = parsed_date.replace(hour=9, minute=0, second=0, microsecond=0)
                            event_end_time = event_start_time.replace(hour=10)
                            event_title = event_title_template.format(date=date)
                            sync_event(event_start_time.isoformat(), event_end_time.isoformat(), event_title)
                        except Exception as e:
                            print(f"Failed to sync event for '{date}': {e}")
        else:
            print(f"Invalid audio file: {audio_file}")
            return "", [], [], []

    except Exception as e:
        print(f"Error during recording or transcription: {str(e).encode('ascii', 'ignore').decode('ascii')}")
        return "", [], [], []

    return transcript, key_points, query_answers, dates

# Main function
async def main():
    transcript = ""
    key_points = []
    query_answers = []
    dates = []

    try:
        transcript, key_points, query_answers, dates = await record_and_transcribe()
        create_word_document(transcript, key_points, query_answers, dates)
    except Exception as e:
        print(f"Error in main process: {str(e).encode('ascii', 'ignore').decode('ascii')}")
        create_word_document("", [], [], [])

if __name__ == "__main__":
    asyncio.run(main())